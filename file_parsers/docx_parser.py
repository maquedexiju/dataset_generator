import os
import re
import io
import base64
from docx import Document
from PIL import Image
from openai import OpenAI

from .basic_parser import BasicParser

img_parse_prompt = '''
你是一个图像识别助手，识别图片中的内容，并返回详细描述，格式如下：
图片概述：图片概述（不超过 30 字）
图片内容：
* 描述 1
* 描述 2
'''

class DOCXParser(BasicParser):

    suffix = 'docx'

    def __init__(self, file_path, root_path, cfg={}, title_prefix='%parent', logger=None, output_dir=''):
        # 检查文件类型
        if not file_path.lower().endswith('.docx'):
            raise ValueError("file type error, not a docx file")

        # 获取图像识别模型的相关算法配置
        self.img_parse_url = cfg['IMG_RECONGNIZE_MODEL']['url']
        self.img_parse_key = cfg['IMG_RECONGNIZE_MODEL']['api_key']
        self.img_parse_model = cfg['IMG_RECONGNIZE_MODEL']['model_name']
        self.openai_client = OpenAI(api_key=self.img_parse_key, base_url=self.img_parse_url)

        super().__init__(file_path, root_path, cfg, title_prefix, logger, output_dir)

        # 检查是否是临时文件
        file_fullname = os.path.basename(file_path)
        if file_fullname.startswith('~') or file_fullname.startswith('.'):
            self.logger.info(f'file {file_fullname} is a temporary file, skip')
            def parse():
                return []
            self.parse = parse

    def _is_list(self, ele):

        if ele.pPr is not None and ele.pPr.numPr is not None and ele.pPr.numPr.ilvl is not None:
            return ele.pPr.numPr.ilvl.val
        else:
            return False


    def save_img(self, img_data, img_path):

        image = Image.open(io.BytesIO(img_data))
        
        # 转换为RGB模式（JPG不支持透明通道）
        if image.mode in ('RGBA', 'LA'):
            background = Image.new('RGB', image.size, (240, 240, 240))
            background.paste(image, mask=image.split()[-1])
            image = background
        
        # 保存为JPG
        image.save(img_path, 'JPEG', quality=95)
        
        
    def get_image_description(self, img_path):
        # 调用 OpenAI 的 API 进行图像识别
        b64_img = base64.b64encode(open(img_path, 'rb').read()).decode()
        response = self.openai_client.chat.completions.create(
            model=self.img_parse_model,
            messages=[
                {"role": "system", "content": img_parse_prompt},
                {"role": "user", "content": [
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/jpeg;base64,{b64_img}"}
                    },
                ]}
            ]
        )
        try:
            result = response.choices[0].message.content
            return result
        except Exception as e:
            self.logger.error(f'image path: {img_path} parse error')
            self.logger.error(f'error: {e}')
            self.logger.error(f'response: {result}')

    def docx_to_markdown(self):
        """
        将 docx 文件转换为 Markdown，处理文本、表格和图片
        :param docx_path: 输入的 docx 文件路径
        :param output_md_path: 输出的 Markdown 文件路径
        :param image_output_dir: 图片保存目录
        """
        doc = self.doc
        markdown_lines = []
        
        # 创建图片输出目录
        image_output_dir = os.path.join(self.output_dir, 'images')
        if not os.path.exists(image_output_dir):
            os.makedirs(image_output_dir)

        # 遍历所有段落和表格
        para_ind = 0
        tbl_ind = 0
        for element in doc.element.body:
            # 处理段落
            if element.tag.endswith('p'):
                paragraph = doc.paragraphs[para_ind]  # 计算段落索引
                para_ind += 1
                markdown_text = self.process_docx_paragraph(paragraph, image_output_dir)
                if markdown_text.strip():  # 避免空行
                    markdown_lines.append(markdown_text)
            
            # 处理表格
            elif element.tag.endswith('tbl'):
                table = doc.tables[tbl_ind]  # 计算表格索引
                tbl_ind += 1
                markdown_table = self.process_docx_table(table)
                markdown_lines.append(markdown_table)
        
        # 写入 Markdown 文件
        md_content = '\n'.join(markdown_lines)
        self.md_content = md_content
        output_md_path = os.path.join(self.output_dir, 'output.md')
        with open(output_md_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        return md_content

    def process_docx_paragraph(self, paragraph, image_output_dir):
        """
        处理段落文本，包括普通文本和图片
        :param paragraph: docx Paragraph 对象
        :param image_output_dir: 图片保存目录
        :return: Markdown 格式的文本
        """
        text = paragraph.text.strip()

        # 根据段落样式添加标题级别
        is_heading = False
        style_name = paragraph.style.name.lower()
        if 'heading' in style_name and text != '':
            level = int(re.search(r'heading (\d+)', style_name).group(1)) if re.search(r'heading (\d+)', style_name) else 1
            text = '#' * level + ' ' + text
            is_heading = True
        # 根据字体大小添加标题级别
        elif self.heading_sizes !=None and text != '':
            s = paragraph.style.font.size
            if s != None and s.pt!= None and s.pt in self.heading_sizes:
                level = self.heading_sizes.index(s.pt) + 1
                text = '#' * level +'' + text
                is_heading = True

        # 检查是否为列表
        is_list = False
        if is_heading is False and text!= '':
            if self._is_list(paragraph._p) is not False:
                # 检查是否为有序列表
                level = paragraph._p.pPr.numPr.ilvl.val
                text = '    ' * (level - 1) + '* ' + text
                is_list = True
            elif self._is_list(paragraph.style.paragraph_format.element) is not False:
                level = paragraph.style.paragraph_format.element.pPr.numPr.ilvl.val
                text ='  '* (level - 1) + '* '+ text
                is_list = True

        # 检查是否有图片
        for run in paragraph.runs:
            if run._element.xpath('.//pic:pic'):
                for shape in run._element.xpath('.//pic:pic'):
                    blip = shape.xpath('.//a:blip/@r:embed')[0]
                    image_part = paragraph.part.related_parts[blip]
                    image_data = image_part.blob
                    
                    # 保存图片
                    image_filename = f"image_{len(os.listdir(image_output_dir)) + 1}.jpg"
                    image_path = os.path.join(image_output_dir, image_filename)
                    self.save_img(image_data, image_path)
                    
                    # 获取图片描述
                    image_description = self.get_image_description(image_path)
                    
                    # 添加图片描述
                    if image_description:
                        img_position = f'{self.knowledge_path}: {image_filename}'
                        text += f"\n\n@resource: {img_position}\n\n{image_description}\n@endresource\n"
        
        if is_list is False:
            text = text + '\n'
        return text

    def process_docx_table(self, table):
        """
        将 docx 表格转换为 Markdown 表格
        :param table: docx Table 对象
        :return: Markdown 格式的表格
        """
        markdown_lines = []
        rows = table.rows
        
        # 表头
        header = [cell.text.strip() for cell in rows[0].cells]
        markdown_lines.append('| ' + ' | '.join(header) + ' |')
        
        # 分隔行
        markdown_lines.append('| ' + ' | '.join(['---'] * len(header)) + ' |')
        
        # 数据行
        for row in rows[1:]:
            row_data = [cell.text.strip() for cell in row.cells]
            markdown_lines.append('| ' + ' | '.join(row_data) + ' |')
        
        return '\n'.join(markdown_lines)


    def split_md(self):
        # 若 md 开头不是一级标题，添加一级标题为文件名
        if not self.md_content.lstrip().startswith('# '):
            self.md_content = f'# {self.file_basename}\n{self.md_content}'

        # 如果 md 中含有多个一级标题，那么 self.title_prefix 保留文件名
        if self.md_content.count('\n# ') > 1:
            if self.file_basename not in self.title_prefix:
                self.title_prefix += '-' + self.file_basename

        result = {}
        headers = []
        current_header_path = ''
        current_content = []

        # 按行遍历 Markdown 内容
        for line in self.md_content.splitlines():
            if line.startswith('#'):
                # 处理当前标题内容
                if current_header_path and current_content and ''.join(current_content) != '':
                    position = f'{self.knowledge_path}: {"-".join(headers)}'
                    result[current_header_path] = f'\n@section: {position}\n\n' + '\n'.join(current_content).strip() + '\n@endsection\n'
                    current_content = []

                # 解析新标题级别和文本
                level = line.count('#')
                header_text = line.strip('#').strip()

                # 标题级别错误处理
                if level > len(headers) + 1:
                    self.logger.warning(f'标题级别跳跃: 当前标题级别 {level}, 上一级标题级别 {len(headers)}。标题: {header_text}')
                    # 尝试修正标题级别，设置为上一级标题级别加 1
                    level = len(headers) + 1

                # 调整标题层级
                if level > len(headers):
                    headers.append(header_text)
                elif level == len(headers):
                    headers[-1] = header_text
                else:
                    headers = headers[:level - 1]
                    headers.append(header_text)

                # current_header_path = '-'.join([self.title_prefix] + headers)
                if self.title_prefix == '':
                    current_header_path = '-'.join(headers)
                else:
                    current_header_path = '-'.join([self.title_prefix] + headers)
            else:
                current_content.append(line)

        # 处理最后一个标题内容
        if current_header_path and current_content:
            position = f'{self.knowledge_path}: {"-".join(headers)}'
            result[current_header_path] = f'\n@section: {position}\n\n' + '\n'.join(current_content).strip() + '\n@endsection\n'
        elif current_content:
            position = self.file_basename
            result[self.file_basename] = self.add_section_tag('\n'.join(current_content).strip(), position)

        self.content_dict = result

    def assemble_qa_info(self):
        qa_info = []

        for title, content in self.content_dict.items():
            qa_info.append({
                'simple_title': title,
                'full_title': title,
                'content': content
            })
        
        self.qa_info = qa_info

    def parse(self):

        try:
            self.doc = Document(self.file_path)
        except Exception as e:
            print(type(e), e)

        self.get_title_level_by_font_size()
        self.docx_to_markdown()
        self.split_md()
        self.assemble_qa_info()

        return self.qa_info

    def get_title_level_by_font_size(self):
        # 逻辑：字数最多的字体大小应该是正文，比这个字体大的就是逐级标题

        heading_sizes = {}
        heading_numbers = 0
        all_text_len = 0
        for p in self.doc.paragraphs:
            # 如果 p 的 style.name 以 Heading 开头，跳过
            if p.style.name.lower().startswith('Heading'):
                heading_numbers += 1
                continue

            s = p.style.font.size
            if s is None: continue

            text_len = len(p.text)
            all_text_len += text_len

            if s.pt not in heading_sizes.keys(): heading_sizes[s.pt] = 0
            heading_sizes[s.pt] += text_len

        # if heading_numbers > 1: # 有标题
        #     self.heading_sizes = None
        # else:

        # 取 value 最大的 heading_sizes
        max_value = max(heading_sizes.values())
        the_most_size = [k for k, v in heading_sizes.items() if v == max_value][0]
        # 取比 the_most_size 大的 key
        bigger_sizes = [k for k in heading_sizes.keys() if k > the_most_size]
        # bigger_sizes 按照 value 倒序
        bigger_sizes = sorted(bigger_sizes, key=lambda x: -heading_sizes[x])
        
        self.heading_sizes = bigger_sizes
